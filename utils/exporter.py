"""
utils/exporter.py — Export conversations to PDF (#10) and Word (#14).
Also exports individual AI answers as standalone PDFs.
"""
import io
from datetime import datetime
from typing import List, Dict, Any


def _clean_for_pdf(text: str) -> str:
    """Strip markdown symbols that PDF can't render."""
    import re
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"#{1,6}\s",      "",    text)
    text = re.sub(r"`(.*?)`",       r"\1", text)
    text = text.replace("<br/>", "\n")
    return text


def export_to_pdf(messages: List[Dict], username: str, module: str) -> bytes:
    """
    Generate a formatted PDF of the conversation.
    Returns raw bytes suitable for st.download_button.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
    )
    from reportlab.lib import colors

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2*cm, bottomMargin=2*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
    )
    styles = getSampleStyleSheet()

    # Custom style definitions
    title_s  = ParagraphStyle("PGTitle",  parent=styles["Title"],
                               fontSize=20, textColor=HexColor("#0d1117"), spaceAfter=4)
    meta_s   = ParagraphStyle("PGMeta",   parent=styles["Normal"],
                               fontSize=9,  textColor=HexColor("#6e7681"), spaceAfter=6)
    qlabel_s = ParagraphStyle("QLabel",   parent=styles["Normal"],
                               fontSize=9,  textColor=HexColor("#4f6ef7"), fontName="Helvetica-Bold")
    qbody_s  = ParagraphStyle("QBody",    parent=styles["Normal"],
                               fontSize=10, textColor=HexColor("#24292f"), spaceAfter=6, leading=14)
    alabel_s = ParagraphStyle("ALabel",   parent=styles["Normal"],
                               fontSize=9,  textColor=HexColor("#10b981"), fontName="Helvetica-Bold")
    abody_s  = ParagraphStyle("ABody",    parent=styles["Normal"],
                               fontSize=10, textColor=HexColor("#24292f"), spaceAfter=4, leading=14)
    src_s    = ParagraphStyle("Src",      parent=styles["Normal"],
                               fontSize=8,  textColor=HexColor("#8b949e"), leftIndent=12)

    story = []

    # Header
    story.append(Paragraph("PayGlobal AI Assistant", title_s))
    story.append(Paragraph("Conversation Export", styles["Heading2"]))
    story.append(Paragraph(
        f"User: <b>{username}</b> &nbsp;|&nbsp; Module: <b>{module}</b> "
        f"&nbsp;|&nbsp; {datetime.now().strftime('%d %B %Y, %H:%M')}",
        meta_s,
    ))
    story.append(HRFlowable(width="100%", thickness=1.5, color=HexColor("#4f6ef7"), spaceAfter=10))

    for msg in messages:
        cleaned = _clean_for_pdf(msg["content"]).replace("\n", "<br/>")

        if msg["role"] == "user":
            story.append(Paragraph("You asked:", qlabel_s))
            story.append(Paragraph(cleaned, qbody_s))
        else:
            story.append(Paragraph("PayGlobal AI:", alabel_s))
            story.append(Paragraph(cleaned, abody_s))
            sources = msg.get("sources", [])
            if sources:
                src_parts = []
                for s in sources[:6]:
                    if isinstance(s, dict):
                        f = s.get("file", "")
                        p = s.get("page", "")
                        src_parts.append(f"{f}{(' — ' + p) if p else ''}")
                    else:
                        src_parts.append(str(s))
                story.append(Paragraph("Sources: " + "  |  ".join(src_parts), src_s))

        story.append(Spacer(1, 0.3*cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#e1e4e8"), spaceAfter=4))
        story.append(Spacer(1, 0.1*cm))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_docx(messages: List[Dict], username: str, module: str) -> bytes:
    """
    Generate a formatted Word document of the conversation.
    Returns raw bytes suitable for st.download_button.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

    # Title
    h = doc.add_heading("PayGlobal AI Assistant", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0d, 0x11, 0x17)

    h2 = doc.add_heading("Conversation Export", 1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(
        f"User: {username}  |  Module: {module}  |  "
        f"Exported: {datetime.now().strftime('%d %B %Y, %H:%M')}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x6e, 0x76, 0x81)

    doc.add_paragraph()  # spacer

    for msg in messages:
        if msg["role"] == "user":
            lbl = doc.add_paragraph()
            r = lbl.add_run("You:")
            r.bold = True
            r.font.color.rgb = RGBColor(0x4f, 0x6e, 0xf7)
            r.font.size = Pt(9)

            body = doc.add_paragraph(msg["content"])
            body.paragraph_format.left_indent = Cm(0.5)
            for run in body.runs:
                run.font.size = Pt(10)

        else:
            lbl = doc.add_paragraph()
            r = lbl.add_run("PayGlobal AI:")
            r.bold = True
            r.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
            r.font.size = Pt(9)

            body = doc.add_paragraph(msg["content"])
            body.paragraph_format.left_indent = Cm(0.5)
            for run in body.runs:
                run.font.size = Pt(10)

            sources = msg.get("sources", [])
            if sources:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(0.5)
                sr = sp.add_run("Sources: ")
                sr.italic = True
                sr.font.size = Pt(8)
                sr.font.color.rgb = RGBColor(0x8b, 0x94, 0x9e)
                parts = []
                for s in sources[:6]:
                    if isinstance(s, dict):
                        f = s.get("file", "")
                        p = s.get("page", "")
                        parts.append(f"{f}{(' — ' + p) if p else ''}")
                    else:
                        parts.append(str(s))
                sp.add_run("  |  ".join(parts)).font.size = Pt(8)

        # Horizontal rule (paragraph border top)
        sep = doc.add_paragraph()
        pPr = sep._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E1E4E8")
        pBdr.append(bottom)
        pPr.append(pBdr)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_answer_pdf(answer: str, sources: list, username: str) -> bytes:
    """
    Generate a standalone PDF of a single AI answer (#10).
    Returns raw bytes suitable for st.download_button.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()

    title_s = ParagraphStyle(
        "AnswerTitle", parent=styles["Title"],
        fontSize=18, textColor=HexColor("#0d1117"), spaceAfter=4,
    )
    meta_s = ParagraphStyle(
        "AnswerMeta", parent=styles["Normal"],
        fontSize=9, textColor=HexColor("#6e7681"), spaceAfter=10,
    )
    body_s = ParagraphStyle(
        "AnswerBody", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#24292f"), leading=15, spaceAfter=6,
    )
    src_s = ParagraphStyle(
        "AnswerSrc", parent=styles["Normal"],
        fontSize=8, textColor=HexColor("#8b949e"), leftIndent=12,
    )

    story = []
    story.append(Paragraph("PayGlobal AI — Answer", title_s))
    story.append(Paragraph(
        f"User: <b>{username}</b> &nbsp;|&nbsp; "
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
        meta_s,
    ))
    story.append(HRFlowable(width="100%", thickness=1.5,
                            color=HexColor("#4f6ef7"), spaceAfter=10))

    cleaned = _clean_for_pdf(answer).replace("\n", "<br/>")
    story.append(Paragraph(cleaned, body_s))

    if sources:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("<b>Sources:</b>", src_s))
        for s in sources[:8]:
            if isinstance(s, dict):
                f = s.get("file", "")
                p = s.get("page", "")
                label = f"{f}{(' — ' + p) if p else ''}"
            else:
                label = str(s)
            story.append(Paragraph(f"&bull; {label}", src_s))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
